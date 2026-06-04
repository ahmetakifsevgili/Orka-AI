from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt


SOURCE = Path(r"D:\Orka\artifacts\bigg-cube\85-numarali-is-fikri-cagri-dokumani.source.doc")
OUT = Path(r"D:\Orka\artifacts\bigg-cube\OrkaOS-BIGG-Cube-Basvuru-Formu.docx")


ANSWERS = {
    "title": "OrkaOS - Kişisel Yapay Zeka Öğrenme İşletim Sistemi",
    "summary": (
        "OrkaOS, öğrencilerin konu öğrenimi, sınav hazırlığı, tekrar, kaynakla çalışma, not üretimi ve "
        "kodlama pratiği gibi dağınık öğrenme süreçlerini tek bir kişisel yapay zeka öğrenme işletim "
        "sisteminde birleştiren global ölçeklenebilir bir EdTech girişimidir. Ürün; Tutor, Mission Control, "
        "Study Coach, Review/Quiz, Source/Wiki, Notebook Studio ve Code Learning IDE modülleriyle öğrencinin "
        "öğrenme kanıtlarını izler, zayıf kavramları ve tekrar ihtiyacını belirler, bir sonraki en doğru "
        "çalışma aksiyonunu önerir. İlk hedef Türkiye'de kontrollü beta/pilot doğrulaması yapmak; ardından "
        "İngilizce arayüz ve çoklu konu desteğiyle uluslararası öğrenci pazarına açılmaktır."
    ),
    "product": (
        "OrkaOS web tabanlı bir kişisel öğrenme işletim sistemidir. Öğrenci sisteme girdiğinde Mission Control "
        "mevcut öğrenme durumunu, öncelikli görevi ve önerilen çalışma modunu gösterir. Tutor açıklama ve telafi "
        "desteği verir; Review/Quiz ölçme ve aralıklı tekrar süreçlerini yürütür; Source/Wiki kaynaklardan güvenli "
        "öğrenme bağlamı üretir; Notebook Studio özet/tekrar paketleri oluşturur; Code Learning IDE programlama "
        "pratiğinde hata ve öğrenme hedefi üzerinden geri bildirim sağlar. Teknik mimari .NET 8 backend, React/"
        "TypeScript frontend, SQL Server, Redis, JWT kimlik doğrulama, API sözleşmeleri ve test otomasyonu üzerine "
        "kuruludur. Gelir modeli B2C abonelik ve ilerleyen aşamada kurum/bootcamp lisanslamasıdır."
    ),
    "problem": (
        "Öğrenciler öğrenme sürecinde sohbet botları, video platformları, not uygulamaları, test bankaları, LMS sistemleri "
        "ve kod çalışma ortamları gibi kopuk araçlar kullanmaktadır. Bu araçlar öğrencinin neyi bildiğini, nerede zorlandığını, "
        "hangi konuyu tekrar etmesi gerektiğini ve bir sonraki doğru adımın ne olduğunu bütünleşik biçimde göstermez. Genel "
        "amaçlı yapay zeka araçları hızlı cevap üretse de öğrencinin geçmiş öğrenme kanıtını, kaynak güvenilirliğini, sınav "
        "hedefini ve tekrar ihtiyacını kalıcı bir öğrenme döngüsüne bağlamakta yetersiz kalır. OrkaOS bu ihtiyacı ortak öğrenme "
        "durumu ve kanıta dayalı aksiyon önerisiyle karşılar."
    ),
    "stage": (
        "İş fikri konsept aşamasını geçmiş, çalışan yazılım prototipi ve kontrollü beta çekirdeği seviyesine ulaşmıştır. "
        "Mevcut sistemde .NET 8 backend, React/TypeScript frontend, SQL Server, Redis, JWT kimlik doğrulama, API sözleşmeleri "
        "ve test katmanları bulunmaktadır. Tutor, Mission Control, Study Coach, Source/Wiki, Notebook Studio, Code Learning IDE "
        "ve quiz/review akışları prototip kapsamında geliştirilmiştir. Mevcut aşama sınırsız üretim kullanımı değil; gerçek "
        "öğrencilerle pilot yapılabilecek kontrollü beta seviyesidir. BİGG desteğiyle UX, canlı pilot, bulut dağıtım ve "
        "ticarileşme hazırlığı tamamlanacaktır."
    ),
    "patent": (
        "İş fikri için şu aşamada yapılmış patent veya faydalı model başvurusu bulunmamaktadır. Mevcut değer; yazılım mimarisi, "
        "öğrenme kanıtı modeli, modüller arası yönlendirme mantığı, güvenli yapay zeka kullanım sınırları, kaynak temelli öğrenme "
        "akışları ve ürün know-how'ı üzerinde oluşmaktadır. Yazılım ürünlerinde patentlenebilirlik sınırları dikkate alınarak "
        "hızlandırma sürecinde marka tescili, kaynak kod/telif koruması, ticari sır niteliğindeki algoritmik akışların dokümantasyonu "
        "ve patentlenebilir teknik bileşen oluşması halinde patent/faydalı model ön araştırması yapılacaktır."
    ),
    "market": (
        "Hedef müşteri ilk aşamada Türkiye'de üniversite öğrencileri, sınavlara hazırlanan bireysel öğrenciler, yazılım öğrenen "
        "gençler ve yoğun kaynak/not yönetimi yapan lise-son sınıf/mezun öğrencilerdir. İlk doğrulama bireysel öğrenci aboneliği "
        "ve küçük pilot gruplarla yapılacaktır. İkinci aşamada hazırlık kursları, bootcamp'ler, üniversite kulüpleri ve eğitim "
        "kurumları için B2B/B2B2C kullanım modeli hedeflenmektedir. Küresel EdTech ve AI in Education pazarı büyümektedir; OrkaOS "
        "bu büyük pazarın tamamını değil, başlangıçta kişisel yapay zeka çalışma kokpiti ihtiyacı olan dar bir giriş segmentini "
        "hedefleyerek pilot metrikleri, kullanıcı tutundurma ve öğrenme aksiyonu tamamlama oranlarıyla ölçeklenmeyi planlamaktadır."
    ),
    "competitors": (
        "Piyasada yerli ve yabancı birçok çözüm bulunmaktadır. ChatGPT, Gemini ve Claude gibi genel amaçlı yapay zeka araçları açıklama "
        "ve soru-cevap desteği sunar. Khanmigo, Quizlet, Duolingo, Coursera, Udemy ve Brilliant belirli öğrenme deneyimlerine odaklanır. "
        "Türkiye'de EBA, Kunduz, Doping Hafıza, Raunt ve çeşitli özel ders/soru çözüm platformları içerik, test veya destek hizmeti verir. "
        "Notion, Google NotebookLM ve LMS sistemleri ise kaynak/not yönetimi için kullanılır. Bu çözümlerin önemli kısmı tek bir öğrenme "
        "moduna odaklanır veya öğrencinin farklı öğrenme kanıtlarını kalıcı, bütünleşik bir kişisel öğrenme durumuna dönüştürmez."
    ),
    "difference": (
        "OrkaOS'un temel farkı, yapay zekayı tekil bir sohbet ekranı olarak değil, öğrencinin öğrenme sürecini yöneten modüler bir işletim "
        "sistemi olarak konumlandırmasıdır. Tutor, Mission Control, Study Coach, Review/Quiz, Source/Wiki, Notebook Studio ve Code Learning "
        "IDE ortak öğrenme durumu etrafında çalışır. Sistem yalnızca cevap üretmez; zayıf kavram, tekrar ihtiyacı, kaynak durumu, sınav "
        "pratiği ve kodlama denemelerini güvenli öğrenme sinyallerine dönüştürür. Kanıt zayıf olduğunda başarı/ustalık iddiası kurmak yerine "
        "tanı, tekrar veya sınırlı telafi aksiyonu önerir. Teknik üstünlük; backend kontrollü yapay zeka kullanımı, kaynak temelli cevaplama, "
        "mahremiyet odaklı veri tasarımı, güvenli fallback ve test edilebilir API kontratlarından gelmektedir."
    ),
    "team": (
        "Proje yürütücüsü Ahmet Akif Sevgili, İstanbul Ticaret Üniversitesi Bilgisayar Mühendisliği mezunudur. OrkaOS'un çekirdek geliştirme "
        "sürecinde yazılım mimarisi, .NET backend, React/TypeScript frontend, SQL Server, Redis, JWT kimlik doğrulama, API sözleşmeleri, test "
        "otomasyonu, yapay zeka entegrasyonları, kaynak temelli öğrenme/RAG ve güvenlik/mahremiyet sınırları üzerinde aktif olarak çalışmıştır. "
        "Mevcut aşamada teknik kurucu olarak prototip geliştirme ve doğrulama çalışmalarını yürütmektedir. Büyüme aşamasında ekip; eğitim "
        "teknolojileri/pedagoji, yapay zeka/veri bilimi, UX/ürün tasarımı, bulut/devops, iş geliştirme ve pilot öğrenci/kurum erişimi alanlarında "
        "tamamlanacaktır. Cube Incubation mentorlukları bu eksik yetkinliklerin tamamlanması ve ticarileşme hazırlığı için kritik görülmektedir."
    ),
    "extra": (
        "BİGG desteğiyle hedef, çalışan prototipi ölçülebilir pilot sonuçlarına ve ticarileşebilir MVP'ye dönüştürmektir. İlk 18 ayda kullanıcı "
        "deneyimi, İngilizce arayüz, bulut dağıtım, KVKK/veri güvenliği, öğrenci pilotları, öğrenme metrikleri, abonelik modeli testi ve kurumsal "
        "iş birlikleri önceliklendirilecektir. Ürün resmi sınav başarısı, puan artışı veya yerleştirme garantisi vermeyecek; değer önerisini kanıta "
        "dayalı öğrenme yönlendirmesi, kaynak dürüstlüğü, tekrar disiplini ve öğrencinin doğru çalışma aksiyonuna yönlendirilmesi üzerine kuracaktır."
    ),
}

LIMITS = {
    "title": 250,
    "summary": 1000,
    "product": 1500,
    "problem": 1500,
    "stage": 1500,
    "patent": 1500,
    "market": 1500,
    "competitors": 1500,
    "difference": 1500,
    "team": 1500,
    "extra": 1000,
}


def set_cell_text(cell, text: str, size: int = 10) -> None:
    cell.text = ""
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(size)


def style_existing_cell(cell, size: int = 10) -> None:
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.name = "Arial"
            run.font.size = Pt(size)


def replace_paragraph_text(paragraph, text: str) -> None:
    paragraph.text = text
    for run in paragraph.runs:
        run.font.name = "Arial"
        run.font.size = Pt(10)


def remove_fixed_table_heights(doc: Document) -> None:
    for table in doc.tables:
        for row in table.rows:
            tr_pr = row._tr.trPr
            if tr_pr is None:
                continue
            for child in list(tr_pr):
                if child.tag in {qn("w:trHeight"), qn("w:cantSplit")}:
                    tr_pr.remove(child)


def main() -> None:
    for key, limit in LIMITS.items():
        length = len(ANSWERS[key])
        if length > limit:
            raise ValueError(f"{key} is {length}/{limit} chars")

    doc = Document(str(SOURCE))

    # Personal information
    set_cell_text(doc.tables[2].rows[1].cells[0], "Ahmet Akif")
    set_cell_text(doc.tables[2].rows[1].cells[2], "Sevgili")
    set_cell_text(doc.tables[3].cell(0, 0), "aakif1345@gmail.com")

    # Main form fields
    set_cell_text(doc.tables[5].cell(0, 0), ANSWERS["title"])
    set_cell_text(doc.tables[6].cell(0, 0), ANSWERS["summary"])
    set_cell_text(doc.tables[7].cell(0, 0), ANSWERS["product"])
    set_cell_text(doc.tables[8].cell(0, 0), ANSWERS["problem"])
    set_cell_text(doc.tables[9].cell(0, 0), ANSWERS["stage"])
    set_cell_text(doc.tables[10].cell(0, 0), ANSWERS["patent"])
    set_cell_text(doc.tables[12].cell(0, 0), ANSWERS["market"])
    set_cell_text(doc.tables[13].cell(0, 0), ANSWERS["competitors"])
    set_cell_text(doc.tables[14].cell(0, 0), ANSWERS["difference"])
    set_cell_text(doc.tables[16].cell(0, 0), ANSWERS["team"])
    set_cell_text(doc.tables[18].cell(0, 0), ANSWERS["extra"])

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                style_existing_cell(cell)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))

    print(f"Wrote: {OUT}")
    for key, limit in LIMITS.items():
        print(f"{key}: {len(ANSWERS[key])}/{limit}")


if __name__ == "__main__":
    main()
